"""
File processor service for extracting product data from various file formats.
Supports: CSV, Excel, PDF, Images, TXT, DOCX
"""
import logging
import io
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

import pandas as pd

logger = logging.getLogger(__name__)


class FileProcessor:
    """
    Process various file formats to extract product data.
    Supports CSV, Excel, PDF, Images, TXT, DOCX.
    """
    
    SUPPORTED_EXTENSIONS = {
        'csv': 'text/csv',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel',
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
    }
    
    # Column name mappings (various possible names -> standardized names)
    COLUMN_MAPPINGS = {
        'nombre': ['nombre', 'name', 'producto', 'product', 'articulo', 'article', 'descripcion', 'description', 'fcarticulo'],
        'precio': ['precio', 'price', 'cost', 'costo', 'valor', 'value', 'fnprecio', 'precio_venta'],
        'tienda_id': ['tienda', 'tienda_id', 'store', 'store_id', 'sucursal', 'fitiendaid', 'fctienda'],
        'codigo': ['codigo', 'code', 'sku', 'barcode', 'upc', 'ean', 'codigo_barras', 'fcarticulo_id'],
        'categoria': ['categoria', 'category', 'tipo', 'type', 'departamento', 'department'],
        'presentacion': ['presentacion', 'size', 'tamaño', 'unidad', 'unit', 'gramaje', 'peso', 'weight'],
    }
    
    def __init__(self):
        logger.info("✅ FileProcessor initialized")
    
    async def process_file(
        self,
        file_data: bytes,
        filename: str,
        content_type: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """
        Process a file and extract product data.
        
        Args:
            file_data: File content as bytes
            filename: Original filename
            content_type: MIME type
            
        Returns:
            Tuple of (list of products, metadata dict)
        """
        extension = Path(filename).suffix.lower().lstrip('.')
        
        logger.info(f"Processing file: {filename} ({content_type})")
        
        try:
            if extension in ('csv',):
                return await self._process_csv(file_data, filename)
            elif extension in ('xlsx', 'xls'):
                return await self._process_excel(file_data, filename)
            elif extension == 'pdf':
                return await self._process_pdf(file_data, filename)
            elif extension in ('txt',):
                return await self._process_txt(file_data, filename)
            elif extension in ('docx', 'doc'):
                return await self._process_docx(file_data, filename)
            elif extension in ('png', 'jpg', 'jpeg'):
                return await self._process_image(file_data, filename)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise
    
    async def _process_csv(
        self, 
        file_data: bytes, 
        filename: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Process CSV file"""
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(io.BytesIO(file_data), encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode CSV file")
        
        return self._dataframe_to_products(df, filename)
    
    async def _process_excel(
        self, 
        file_data: bytes, 
        filename: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Process Excel file"""
        df = pd.read_excel(io.BytesIO(file_data))
        return self._dataframe_to_products(df, filename)
    
    async def _process_pdf(
        self, 
        file_data: bytes, 
        filename: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Process PDF file - extract tables or text"""
        try:
            import pdfplumber
        except ImportError:
            raise ValueError("PDF processing requires pdfplumber library")
        
        products = []
        all_text = []
        
        with pdfplumber.open(io.BytesIO(file_data)) as pdf:
            for page in pdf.pages:
                # Try to extract tables first
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:
                        # Assume first row is header
                        headers = [str(h).lower().strip() if h else '' for h in table[0]]
                        for row in table[1:]:
                            if row and any(cell for cell in row):
                                product = self._map_row_to_product(headers, row)
                                if product:
                                    products.append(product)
                
                # Also extract text
                text = page.extract_text()
                if text:
                    all_text.append(text)
        
        # If no tables found, try to parse text
        if not products and all_text:
            products = self._parse_text_for_products("\n".join(all_text))
        
        metadata = {
            "filename": filename,
            "format": "pdf",
            "total_products": len(products),
            "pages_processed": len(all_text) if all_text else 0
        }
        
        return products, metadata
    
    async def _process_txt(
        self, 
        file_data: bytes, 
        filename: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Process TXT file"""
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_data.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode TXT file")
        
        # Try to detect if it's a tab/comma separated file
        lines = text.strip().split('\n')
        
        if len(lines) > 1:
            # Check for delimiters
            first_line = lines[0]
            if '\t' in first_line:
                # Tab separated
                df = pd.read_csv(io.StringIO(text), sep='\t')
                return self._dataframe_to_products(df, filename)
            elif ',' in first_line and first_line.count(',') > 1:
                # Comma separated
                df = pd.read_csv(io.StringIO(text))
                return self._dataframe_to_products(df, filename)
        
        # Parse as free text
        products = self._parse_text_for_products(text)
        
        metadata = {
            "filename": filename,
            "format": "txt",
            "total_products": len(products)
        }
        
        return products, metadata
    
    async def _process_docx(
        self, 
        file_data: bytes, 
        filename: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Process DOCX file"""
        try:
            from docx import Document
        except ImportError:
            raise ValueError("DOCX processing requires python-docx library")
        
        doc = Document(io.BytesIO(file_data))
        products = []
        
        # Extract tables
        for table in doc.tables:
            if len(table.rows) > 1:
                headers = [cell.text.lower().strip() for cell in table.rows[0].cells]
                for row in table.rows[1:]:
                    values = [cell.text.strip() for cell in row.cells]
                    product = self._map_row_to_product(headers, values)
                    if product:
                        products.append(product)
        
        # If no tables, extract paragraphs
        if not products:
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            products = self._parse_text_for_products(text)
        
        metadata = {
            "filename": filename,
            "format": "docx",
            "total_products": len(products)
        }
        
        return products, metadata
    
    async def _process_image(
        self, 
        file_data: bytes, 
        filename: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """
        Process image file using OCR.
        Requires external OCR service or Gemini Vision.
        """
        # For now, return empty - this would need OCR integration
        logger.warning(f"Image processing not fully implemented: {filename}")
        
        metadata = {
            "filename": filename,
            "format": "image",
            "total_products": 0,
            "note": "Image processing requires OCR integration"
        }
        
        return [], metadata
    
    def _dataframe_to_products(
        self, 
        df: pd.DataFrame, 
        filename: str
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Convert DataFrame to list of product dicts"""
        # Normalize column names
        df.columns = [str(c).lower().strip() for c in df.columns]
        
        # Map columns
        column_map = {}
        for standard_name, possible_names in self.COLUMN_MAPPINGS.items():
            for col in df.columns:
                if col in possible_names:
                    column_map[col] = standard_name
                    break
        
        # Check required columns
        mapped_cols = set(column_map.values())
        if 'nombre' not in mapped_cols:
            # Try to find any text column
            for col in df.columns:
                if df[col].dtype == 'object':
                    column_map[col] = 'nombre'
                    break
        
        if 'nombre' not in set(column_map.values()):
            raise ValueError("Could not find product name column in file")
        
        # Rename columns
        df = df.rename(columns=column_map)
        
        # Convert to products
        products = []
        for _, row in df.iterrows():
            product = {}
            
            # Required: nombre
            nombre = row.get('nombre', '')
            if pd.isna(nombre) or not str(nombre).strip():
                continue
            product['nombre'] = str(nombre).strip().upper()
            
            # Optional: precio
            precio = row.get('precio', 0)
            if pd.notna(precio):
                try:
                    product['precio'] = float(str(precio).replace('$', '').replace(',', '').strip())
                except ValueError:
                    product['precio'] = 0
            else:
                product['precio'] = 0
            
            # Optional: tienda_id
            tienda = row.get('tienda_id', '')
            if pd.notna(tienda):
                product['tienda_id'] = str(tienda).strip()
            
            # Optional: codigo
            codigo = row.get('codigo', '')
            if pd.notna(codigo):
                product['codigo'] = str(codigo).strip()
            
            # Optional: categoria
            categoria = row.get('categoria', '')
            if pd.notna(categoria):
                product['categoria'] = str(categoria).strip()
            
            # Optional: presentacion
            presentacion = row.get('presentacion', '')
            if pd.notna(presentacion):
                product['presentacion'] = str(presentacion).strip()
            
            products.append(product)
        
        metadata = {
            "filename": filename,
            "format": "tabular",
            "total_products": len(products),
            "columns_found": list(df.columns),
            "columns_mapped": column_map
        }
        
        return products, metadata
    
    def _map_row_to_product(
        self, 
        headers: List[str], 
        values: List[str]
    ) -> Optional[Dict[str, Any]]:
        """Map a table row to a product dict"""
        if len(headers) != len(values):
            return None
        
        row_dict = dict(zip(headers, values))
        
        # Map columns
        product = {}
        for standard_name, possible_names in self.COLUMN_MAPPINGS.items():
            for col_name in possible_names:
                if col_name in row_dict and row_dict[col_name]:
                    value = str(row_dict[col_name]).strip()
                    if value:
                        if standard_name == 'precio':
                            try:
                                product[standard_name] = float(value.replace('$', '').replace(',', ''))
                            except ValueError:
                                pass
                        else:
                            product[standard_name] = value.upper() if standard_name == 'nombre' else value
                        break
        
        return product if product.get('nombre') else None
    
    def _parse_text_for_products(self, text: str) -> List[Dict[str, Any]]:
        """
        Parse free text to extract product information.
        Looks for patterns like "Product Name - $XX.XX" or "Product Name: XX.XX"
        """
        products = []
        
        # Pattern: Product name followed by price
        patterns = [
            r'(.+?)\s*[-:]\s*\$?(\d+\.?\d*)',
            r'(.+?)\s+\$(\d+\.?\d*)',
        ]
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            for pattern in patterns:
                match = re.search(pattern, line)
                if match:
                    nombre = match.group(1).strip().upper()
                    try:
                        precio = float(match.group(2))
                    except ValueError:
                        precio = 0
                    
                    if nombre and len(nombre) > 2:
                        products.append({
                            'nombre': nombre,
                            'precio': precio
                        })
                    break
        
        return products
    
    def get_preview_data(
        self, 
        products: List[Dict[str, Any]], 
        limit: int = 100
    ) -> List[Dict[str, Any]]:
        """Get preview data for display"""
        return products[:limit]
